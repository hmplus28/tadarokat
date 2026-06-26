"""Import داده‌های گردش کار (تحویل، دستور، استعلام و ...) از اکسل — برای داده‌های قبلی."""

from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List, Tuple

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

from services import history_service, local_storage
from services.excel_service import invalidate_cache
from services.local_storage import LOCAL_SHEETS, SHEET_HEADERS

IMPORT_ORDER = (
    "issued_inquiries",
    "pre_invoices",
    "pre_invoice_lines",
    "orders",
    "deliveries",
    "purchase_edits",
    "product_history",
)

SHEET_ALIASES: Dict[str, str] = {
    "راهنما": "_guide",
    "guide": "_guide",
    "استعلام صادر شده": "issued_inquiries",
    "issued_inquiries": "issued_inquiries",
    "پیش فاکتور پیمانکار": "pre_invoices",
    "پیش فاکتور": "pre_invoices",
    "pre_invoices": "pre_invoices",
    "ردیف پیش فاکتور": "pre_invoice_lines",
    "pre_invoice_lines": "pre_invoice_lines",
    "دستور خرید": "orders",
    "دستور خرید محلی": "orders",
    "orders": "orders",
    "تحویل": "deliveries",
    "تحویل محلی": "deliveries",
    "deliveries": "deliveries",
    "ویرایش درخواست": "purchase_edits",
    "purchase_edits": "purchase_edits",
    "سابقه خرید کالا": "product_history",
    "سابقه خرید": "product_history",
    "product_history": "product_history",
}

TEMPLATE_SAMPLES: Dict[str, Dict[str, Any]] = {
    "issued_inquiries": {
        "شماره استعلام": "123456",
        "شماره درخواست خرید": "1001",
        "نوع خرید": "عادی",
        "تاریخ استعلام": "1404/01/15",
        "وضعیت": "تایید شده",
        "مهلت استعلام": "1404/02/01",
        "انبار": "انبار مصرفی",
        "درخواست دهنده": "واحد تولید",
        "کارشناس خرید": "مصطفی رضوانی",
        "صادر کننده سند": "مصطفی رضوانی",
    },
    "pre_invoices": {
        "id": "",
        "شماره استعلام": "123456",
        "شماره پیش فاکتور": "PF-1001",
        "نام پیمانکار": "شرکت نمونه",
        "شهر پیمانکار": "تهران",
        "تاریخ پیش فاکتور": "1404/01/16",
        "نوع فاکتور": "رسمی",
        "جمع کل": "11000000",
        "وضعیت مدیر": "تایید شده",
    },
    "pre_invoice_lines": {
        "id": "",
        "preinvoice_id": "(id پیش‌فاکتور)",
        "ردیف": "1",
        "عنوان کالا": "پیچ M10",
        "واحد": "عدد",
        "فی": "50000",
        "تعداد": "100",
        "جمع کل": "5000000",
    },
    "orders": {
        "id": "",
        "شماره دستور": "DH-1404-001",
        "شماره استعلام": "123456",
        "شماره خرید": "1001",
        "ردیف": "1",
        "عنوان کالا": "پیچ M10",
        "انبار": "انبار مصرفی",
        "کارشناس": "مصطفی رضوانی",
        "نام پیمانکار": "شرکت نمونه",
        "مرحله فعلی": "تحویل شده",
        "تاریخ دستور": "1404/01/20",
    },
    "deliveries": {
        "id": "",
        "شماره تحویل": "TH-1404-001",
        "شماره دستور": "DH-1404-001",
        "شماره خرید": "1001",
        "عنوان کالا": "پیچ M10",
        "انبار": "انبار مصرفی",
        "مقدار": "100",
        "واحد": "عدد",
        "تاریخ تحویل": "1404/02/05",
        "تحویل گیرنده": "انباردار",
        "وضعیت": "تحویل شده",
    },
}


def _resolve_sheet_key(name: str):
    key = str(name or "").strip()
    if key in SHEET_ALIASES:
        resolved = SHEET_ALIASES[key]
        return None if resolved == "_guide" else resolved
    for alias, sk in SHEET_ALIASES.items():
        if alias in key:
            return None if sk == "_guide" else sk
    return None


def _is_template_sample(rec: Dict[str, Any], sheet_key: str) -> bool:
    sample = TEMPLATE_SAMPLES.get(sheet_key, {})
    if not sample:
        return False
    matches = 0
    checked = 0
    for k, v in sample.items():
        if v in (None, "", "(id پیش‌فاکتور)"):
            continue
        checked += 1
        if str(rec.get(k, "")).strip() == str(v).strip():
            matches += 1
    return checked >= 2 and matches >= checked - 1


def _df_to_records(df: pd.DataFrame) -> List[Dict[str, Any]]:
    if df.empty:
        return []
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    df = df.dropna(how="all")
    records = []
    for _, row in df.iterrows():
        rec = {}
        for k, v in row.items():
            if pd.isna(v):
                continue
            s = str(v).strip()
            if s and s.lower() not in ("nan", "none"):
                rec[k] = v
        if rec:
            records.append(rec)
    return records


def build_excel_template() -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "راهنما"
    guide = [
        ["راهنمای import داده گردش کار — سامانه تدارکات"],
        [""],
        ["۱. هر شیت = یک نوع داده (استعلام، پیش‌فاکتور، دستور، تحویل و ...)"],
        ["۲. سطر اول = نام ستون‌ها — سطرهای بعد = داده"],
        ["۳. برای به‌روزرسانی: همان کلید یکتا را وارد کنید (مثلاً شماره استعلام یا id)"],
        ["۴. id خالی = ردیف جدید"],
        ["۵. ترتیب پیشنهادی: استعلام → پیش‌فاکتور → ردیف پیش‌فاکتور → دستور → تحویل"],
        [""],
        ["شیت‌های موجود در این فایل:"],
    ]
    for sk in IMPORT_ORDER:
        guide.append([LOCAL_SHEETS.get(sk, sk)])
    for r_idx, row in enumerate(guide, 1):
        for c_idx, val in enumerate(row, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            if r_idx == 1:
                cell.font = Font(bold=True, size=12)
    ws.column_dimensions["A"].width = 72

    header_fill = PatternFill("solid", fgColor="EEF2FF")
    header_font = Font(bold=True, color="312E81")
    for sk in IMPORT_ORDER:
        title = LOCAL_SHEETS.get(sk, sk)[:31]
        sheet = wb.create_sheet(title)
        headers = SHEET_HEADERS[sk]
        sample = TEMPLATE_SAMPLES.get(sk, {})
        for col, h in enumerate(headers, 1):
            cell = sheet.cell(row=1, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center")
            sheet.cell(row=2, column=col, value=sample.get(h, ""))

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_word_guide() -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    title = doc.add_heading("راهنمای ورود داده‌های گردش کار", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = intro.add_run(
        "این سند نحوهٔ وارد کردن داده‌های انجام‌شده (استعلام، دستور خرید، تحویل و ...) "
        "به سامانه تدارکات را توضیح می‌دهد. داده‌ها از طریق فایل اکسل تمپلیت وارد می‌شوند."
    )
    run.font.size = Pt(11)

    doc.add_heading("مراحل کار", level=1)
    steps = [
        "از پنل مدیر سیستم، دکمه «دانلود تمپلیت اکسل» را بزنید.",
        "در هر شیت، ستون‌ها را پر کنید. سطر نمونه را می‌توانید حذف یا جایگزین کنید.",
        "برای تحویل‌های قبلی: شیت «تحویل محلی» را پر کنید.",
        "برای دستورات صادرشده: شیت «دستور خرید محلی».",
        "فایل را در پنل سیستم با «انتخاب فایل» و «اجرای Import» بارگذاری کنید.",
        "پس از import، داده‌ها در بخش‌های مربوط (تحویل‌ها، دستور خرید، استعلام‌ها) نمایش داده می‌شوند.",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(s)

    doc.add_heading("شیت‌های اکسل", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "نام شیت"
    hdr[1].text = "کاربرد"
    rows_data = [
        ("استعلام صادر شده", "استعلام‌های صادر شده قبلی"),
        ("پیش فاکتور پیمانکار", "پیش‌فاکتور پیمانکاران"),
        ("ردیف پیش فاکتور", "اقلام هر پیش‌فاکتور"),
        ("دستور خرید محلی", "دستورات خرید صادرشده"),
        ("تحویل محلی", "تحویل‌های انجام‌شده"),
        ("ویرایش درخواست", "اصلاحات روی درخواست خرید"),
        ("سابقه خرید کالا", "سابقه قیمت خرید"),
    ]
    for name, use in rows_data:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = use

    doc.add_heading("نکات مهم", level=1)
    notes = [
        "کلید یکتا: استعلام = شماره استعلام؛ دستور/تحویل = id یا شماره دستور/تحویل.",
        "اگر id خالی باشد، ردیف جدید ساخته می‌شود.",
        "اگر کلید تکراری باشد، ردیف موجود به‌روز می‌شود.",
        "تاریخ‌ها را به فرمت شمسی وارد کنید (مثال: 1404/01/15).",
        "پس از import، سوپر‌یوزر می‌تواند هر فیلد را با آیکون مداد ویرایش کند.",
    ]
    for n in notes:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(n)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("سامانه تدارکات — تمپلیت ورود داده")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def import_from_excel(content: bytes, username: str) -> Dict[str, Any]:
    try:
        xls = pd.ExcelFile(BytesIO(content), engine="openpyxl")
    except Exception as exc:
        raise ValueError(f"فایل اکسل نامعتبر است: {exc}") from exc

    sheet_results: Dict[str, Dict[str, int]] = {}
    parsed: List[Tuple[str, List[Dict[str, Any]]]] = []

    for sheet_name in xls.sheet_names:
        sk = _resolve_sheet_key(sheet_name)
        if not sk or sk not in IMPORT_ORDER:
            continue
        df = pd.read_excel(xls, sheet_name=sheet_name, engine="openpyxl")
        records = [r for r in _df_to_records(df) if not _is_template_sample(r, sk)]
        if records:
            parsed.append((sk, records))

    if not parsed:
        raise ValueError("هیچ شیت معتبری در فایل یافت نشد — نام شیت‌ها را با تمپلیت مقایسه کنید")

    order_map = {sk: i for i, sk in enumerate(IMPORT_ORDER)}
    parsed.sort(key=lambda x: order_map.get(x[0], 99))

    total_ins = total_upd = total_skip = 0
    for sk, records in parsed:
        stats = local_storage.upsert_sheet_rows(sk, records, username)
        sheet_results[LOCAL_SHEETS.get(sk, sk)] = stats
        total_ins += stats["inserted"]
        total_upd += stats["updated"]
        total_skip += stats["skipped"]
        history_service.log_action(
            "import گردش کار", sk, username, "import",
            f"+{stats['inserted']} / ~{stats['updated']}",
        )

    invalidate_cache()
    return {
        "ok": True,
        "sheets": sheet_results,
        "totals": {"inserted": total_ins, "updated": total_upd, "skipped": total_skip},
        "message": f"Import انجام شد: {total_ins} جدید، {total_upd} به‌روز، {total_skip} رد شده",
    }